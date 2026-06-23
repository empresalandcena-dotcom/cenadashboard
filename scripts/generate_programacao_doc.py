from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Programacao_Producao_Diaria_Plano.docx"


def set_run_font(run, name="Arial", size=11, bold=False, color="000000"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_para_spacing(paragraph, before=0, after=8, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=20)
        set_para_spacing(p, before=20, after=6)
    elif level == 2:
        set_run_font(run, size=16)
        set_para_spacing(p, before=18, after=6)
    else:
        set_run_font(run, size=14, color="434343")
        set_para_spacing(p, before=16, after=4)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=11)
    set_para_spacing(p, before=0, after=8)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11)
        set_para_spacing(p, before=0, after=4)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="000000"):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_font(run, size=10, bold=bold, color=color)
    set_para_spacing(p, before=0, after=0, line=1.0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(2.2), Inches(2.0), Inches(2.3)]

    hdr = table.rows[0].cells
    headers = ["Bloco", "Objetivo", "Entregavel"]
    for idx, cell in enumerate(hdr):
        cell.width = widths[idx]
        shade_cell(cell, "163847")
        set_cell_text(cell, headers[idx], bold=True, color="FFFFFF")

    rows = [
        ("Base de dados", "Padronizar a planilha principal", "Tabela confiavel para leitura automatica"),
        ("Indicadores", "Definir regras unicas de calculo", "KPIs reconciliados entre cards e tabelas"),
        ("Integracao", "Ler Google Sheets ou Excel Online", "Atualizacao quase em tempo real"),
        ("Governanca", "Definir rotina e responsavel", "Uso gerencial com confianca"),
    ]

    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].width = widths[idx]
            set_cell_text(cells[idx], value)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    title = doc.add_paragraph()
    title_run = title.add_run("Plano Inicial - Programacao e Producao Diaria")
    set_run_font(title_run, size=26)
    set_para_spacing(title, before=0, after=3, line=1.0)

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run(
        "Documento de orientacao para conectar dados reais ao dashboard da CENA com consistencia operacional."
    )
    set_run_font(sub_run, size=11, color="555555")
    set_para_spacing(subtitle, before=0, after=10)

    add_heading(doc, "Objetivo", 1)
    add_body(
        doc,
        "Organizar a frente de Programacao / Producao Diaria para que o dashboard use dados reais, indicadores confiaveis e uma rotina de atualizacao sustentavel.",
    )

    add_heading(doc, "Perguntas que o painel precisa responder", 1)
    add_bullets(
        doc,
        [
            "Quanto foi programado no periodo.",
            "Quanto foi concluido no periodo.",
            "Quais equipes, carteiras e municipios estao acima ou abaixo da meta.",
            "Onde estao os principais gargalos entre programacao, reprogramacao, cancelamento e conclusao.",
        ],
    )

    add_heading(doc, "Base recomendada", 1)
    add_bullets(
        doc,
        [
            "Fonte principal: 06-PROGRAMACAO.xlsx.",
            "Complemento recomendado: 02_CONTROLE_STATUS.xlsx.",
            "Camada online oficial: Google Sheets ou Excel Online.",
        ],
    )

    add_heading(doc, "Colunas minimas da estrutura", 1)
    add_bullets(
        doc,
        [
            "nota, pep, municipio, carteira, equipe.",
            "status_programacao, status_pi, data_programada, data_prevista, data_conclusao.",
            "semana_programacao, mes_programacao, qtd_postes, qtd_medidores.",
            "valor_mao_obra, valor_medido, valor_faturado, ultima_atualizacao.",
        ],
    )

    add_heading(doc, "Indicadores prioritarios", 1)
    add_heading(doc, "KPIs executivos", 2)
    add_bullets(
        doc,
        [
            "Programadas no periodo.",
            "Concluidas no periodo.",
            "Reprogramadas e canceladas.",
            "Taxa de conclusao.",
            "Taxa de reprogramacao.",
            "Valor programado, concluido e em aberto.",
        ],
    )

    add_heading(doc, "KPIs gerenciais", 2)
    add_bullets(
        doc,
        [
            "Produtividade por equipe, carteira, municipio e semana.",
            "Media diaria de producao.",
            "Atingimento de meta por equipe.",
            "Backlog operacional.",
            "Aging medio das notas em aberto.",
        ],
    )

    add_heading(doc, "Regras de calculo", 1)
    add_bullets(
        doc,
        [
            "Programadas = quantidade de notas com data programada no periodo filtrado.",
            "Concluidas = quantidade de notas com status concluido.",
            "Taxa de conclusao = concluidas / programadas.",
            "Backlog operacional = programadas e nao concluidas.",
            "Atingimento da meta = realizado / meta definida para o periodo.",
        ],
    )

    add_heading(doc, "Sequencia de implantacao", 1)
    build_table(doc)

    add_heading(doc, "Checklist de pronto", 1)
    add_bullets(
        doc,
        [
            "Base sem duplicidade critica.",
            "Status padronizados.",
            "Datas validas e campos numericos consistentes.",
            "Indicadores reconciliados entre cards, graficos e tabela.",
            "Atualizacao online definida com responsavel claro.",
        ],
    )

    add_heading(doc, "Proximo passo recomendado", 1)
    add_body(
        doc,
        "O melhor passo agora e mapear o arquivo 06-PROGRAMACAO.xlsx: abas usadas, colunas obrigatorias, campos com problema de padronizacao e indicadores que podem nascer direto da base. Com isso, a ligacao da pagina aos dados reais fica muito mais segura.",
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
